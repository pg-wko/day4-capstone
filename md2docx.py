"""Convert the capstone playbook Markdown into a formatted .docx."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

CODE_FONT = "Consolas"
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\s][^*]*?\*)")


def shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def add_left_bar(paragraph, color="4472C4"):
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    paragraph._p.get_or_add_pPr().append(borders)


def add_inline(paragraph, text):
    """Render **bold**, *italic* and `code` spans into a paragraph."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x28, 0x28)
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)


def add_code_block(doc, lines, language):
    if language == "mermaid":
        note = doc.add_paragraph()
        run = note.add_run(
            "Diagram (Mermaid source \u2014 paste into mermaid.live or redraw before submitting)"
        )
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        note.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), "F4F4F4")
    cell.text = ""
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line)
        run.font.name = CODE_FONT
        run.font.size = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        para = cell.paragraphs[0]
        add_inline(para, text)
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
        shade(cell._tc.get_or_add_tcPr(), "D9E2F3")

    for row in body:
        cells = table.add_row().cells
        for cell, text in zip(cells, row[: len(header)]):
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline(para, text)
            for run in para.runs:
                run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.8)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip().lower()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block, language)
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_separator(lines[i]):
                    rows.append(split_row(lines[i]))
                i += 1
            if rows:
                width = max(len(r) for r in rows)
                rows = [r + [""] * (width - len(r)) for r in rows]
                add_table(doc, rows)
            continue

        if re.fullmatch(r"-{3,}", stripped):
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            i += 1
            continue

        heading = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).replace("**", "")
            if level == 1:
                para = doc.add_heading(text, level=0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level=min(level - 1, 4))
            i += 1
            continue

        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            add_inline(para, " ".join(t for t in block if t))
            for run in para.runs:
                run.font.size = Pt(10)
            add_left_bar(para)
            continue

        bullet = re.match(r"[-*]\s+(.*)", stripped)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"\d+\.\s+(.*)", stripped)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            add_inline(para, numbered.group(1))
            i += 1
            continue

        para = doc.add_paragraph()
        add_inline(para, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    convert(src, dst)
